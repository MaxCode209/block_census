"""API routes for the application."""
import io
import re
import zipfile
from decimal import Decimal

from flask import Blueprint, jsonify, request, send_file
from sqlalchemy.orm import Session, load_only
from sqlalchemy import or_, text
from typing import List, Dict, Optional, Tuple, Any

from backend.database import get_db
from backend.models import CensusData, SchoolData, School, AttendanceZone, CountyEmployer

# Columns that exist in census_data table (no city until added in Supabase)
_CENSUS_LOAD_COLUMNS = (
    CensusData.id, CensusData.zip_code, CensusData.county,
    CensusData.population, CensusData.median_age, CensusData.average_household_income,
    CensusData.data_year, CensusData.created_at, CensusData.updated_at,
)
# Explicit column list for raw SQL (includes school counts and ratings)
_CENSUS_SQL_COLS = "id, zip_code, county, population, median_age, average_household_income, local_employment_rating, data_year, created_at, updated_at, total_schools, elementary_schools, middle_schools, high_schools, average_school_rating, average_elementary_school_rating, average_middle_school_rating, average_high_school_rating"
from backend.census_api import CensusAPIClient
from backend.zone_utils import (
    point_in_polygon,
    find_zoned_schools,
    find_all_zoned_schools,
    load_zip_polygon,
    zones_intersecting_zip,
    zones_intersecting_zip_diagnostic,
    group_zones_by_district,
    district_geometry_in_zip,
    zone_geometry_in_zip,
)
from backend.greatschools_client import GreatSchoolsClient

api = Blueprint('api', __name__, url_prefix='/api')

# Cache for school-zones by_level responses (zip -> response dict). Max 50 zips.
_school_zones_cache = {}
_SCHOOL_ZONES_CACHE_MAX = 50

def _census_kwargs(data):
    """Filter dict to only keys that exist on CensusData model (avoids extra columns from API)."""
    allowed = {c.key for c in CensusData.__table__.c}
    return {k: v for k, v in (data or {}).items() if k in allowed}


@api.route('/employers', methods=['GET'])
def get_employers():
    """Get county employers; optional filter by county. Lightweight endpoint for data validation."""
    try:
        db: Session = next(get_db())
    except Exception as e:
        return jsonify({'error': f'Database connection failed: {str(e)}', 'data': []}), 500
    try:
        county = request.args.get('county', type=str)
        limit = request.args.get('limit', type=int, default=500)
        q = db.query(CountyEmployer).order_by(CountyEmployer.county_name, CountyEmployer.rank)
        if county and county.strip():
            q = q.filter(CountyEmployer.county_name.ilike(f"%{county.strip()}%"))
        rows = q.limit(limit).all()
        return jsonify({'data': [r.to_dict() for r in rows]})
    except Exception as e:
        return jsonify({'error': str(e), 'data': []}), 500
    finally:
        db.close()


@api.route('/census-data', methods=['GET'])
def get_census_data():
    """Get census data with optional filters. Uses raw SQL so we never reference city column."""
    try:
        db: Session = next(get_db())
    except Exception as e:
        return jsonify({'error': f'Database connection failed: {str(e)}', 'data': []}), 500

    try:
        zip_code = request.args.get('zip_code')
        city = request.args.get('city')
        state = request.args.get('state')  # Optional: e.g. "NC" to disambiguate Wilmington
        min_income = request.args.get('min_income', type=float)
        max_income = request.args.get('max_income', type=float)
        min_population = request.args.get('min_population', type=int)
        max_population = request.args.get('max_population', type=int)
        min_age = request.args.get('min_age', type=float)
        max_age = request.args.get('max_age', type=float)
        min_employment_rating = request.args.get('min_employment_rating', type=float)
        min_elementary_school_rating = request.args.get('min_elementary_school_rating', type=float)
        min_blended_school_rating = request.args.get('min_blended_school_rating', type=float)
        limit = request.args.get('limit', type=int, default=1000)
        offset = request.args.get('offset', type=int, default=0)

        # Build WHERE and params (state filter uses census_data.state when column exists)
        use_state_filter = state and str(state).strip()
        t = ""  # No table alias

        where_parts = []
        params = {}
        if zip_code:
            where_parts.append("zip_code = :zip_code")
            params["zip_code"] = zip_code
        if city:
            where_parts.append("LOWER(TRIM(COALESCE(city, ''))) = LOWER(TRIM(:city))")
            params["city"] = city.strip()
        if use_state_filter:
            where_parts.append("UPPER(TRIM(COALESCE(state, ''))) = UPPER(TRIM(:state))")
            params["state"] = str(state).strip()
        if min_income:
            where_parts.append("average_household_income >= :min_income")
            params["min_income"] = min_income
        if max_income:
            where_parts.append("average_household_income <= :max_income")
            params["max_income"] = max_income
        if min_population:
            where_parts.append("population >= :min_population")
            params["min_population"] = min_population
        if max_population:
            where_parts.append("population <= :max_population")
            params["max_population"] = max_population
        if min_age is not None:
            where_parts.append("median_age >= :min_age")
            params["min_age"] = min_age
        if max_age is not None:
            where_parts.append("median_age <= :max_age")
            params["max_age"] = max_age
        if min_employment_rating is not None:
            where_parts.append("local_employment_rating IS NOT NULL AND local_employment_rating >= :min_employment_rating")
            params["min_employment_rating"] = min_employment_rating
        if min_elementary_school_rating is not None:
            where_parts.append("average_elementary_school_rating IS NOT NULL AND average_elementary_school_rating >= :min_elementary_school_rating")
            params["min_elementary_school_rating"] = min_elementary_school_rating
        if min_blended_school_rating is not None:
            where_parts.append("average_school_rating IS NOT NULL AND average_school_rating >= :min_blended_school_rating")
            params["min_blended_school_rating"] = min_blended_school_rating
        where_sql = " AND ".join(where_parts) if where_parts else "1=1"

        from_clause = "census_data"
        col_list = _CENSUS_SQL_COLS
        keys = ["id", "zip_code", "county", "population", "median_age", "average_household_income", "local_employment_rating", "data_year", "created_at", "updated_at", "total_schools", "elementary_schools", "middle_schools", "high_schools", "average_school_rating", "average_elementary_school_rating", "average_middle_school_rating", "average_high_school_rating"]

        # Count with raw SQL
        count_sql = text(f"SELECT COUNT(*) FROM {from_clause} WHERE {where_sql}")
        total = db.execute(count_sql, params).scalar()

        # Data with raw SQL
        order_col = "zip_code"
        data_sql = text(
            f"SELECT {col_list} FROM {from_clause} WHERE {where_sql} "
            f"ORDER BY {order_col} LIMIT :lim OFFSET :off"
        )
        params["lim"] = limit
        params["off"] = offset
        rows = db.execute(data_sql, params).fetchall()

        # Build response dicts (same shape as to_dict)
        data = []
        for row in rows:
            d = dict(zip(keys, row))
            if d.get("local_employment_rating") is not None:
                d["local_employment_rating"] = float(d["local_employment_rating"])
            if d.get("created_at"):
                d["created_at"] = d["created_at"].isoformat() if hasattr(d["created_at"], "isoformat") else str(d["created_at"])
            if d.get("updated_at"):
                d["updated_at"] = d["updated_at"].isoformat() if hasattr(d["updated_at"], "isoformat") else str(d["updated_at"]) if d["updated_at"] else None
            data.append(d)

        return jsonify({
            "data": data,
            "total": total,
            "limit": limit,
            "offset": offset,
        })
    except Exception as e:
        return jsonify({"error": str(e), "data": []}), 500

@api.route('/census-block-groups', methods=['GET'])
def get_census_block_groups():
    """
    Get census block group data for map display.
    Supports: search by city, search by address/zip (lat/lng or zip_code).
    Returns block groups with geometry (GeoJSON) for drawing boundaries.
    If format=shapefile, returns LandVision-ready ZIP instead (same endpoint, so no 404).
    If format=test-zip, returns a minimal valid ZIP to verify downloads work.
    """
    if request.args.get('format') == 'test-zip':
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_STORED) as zf:
            zf.writestr('hello.txt', b'LandVision export test - if you see this, download is OK')
        zip_bytes = zip_buf.getvalue()
        from flask import Response
        r = Response(zip_bytes, mimetype='application/zip')
        r.headers['Content-Disposition'] = 'attachment; filename="test_download.zip"'
        r.headers['Content-Length'] = len(zip_bytes)
        return r
    if request.args.get('format') == 'shapefile':
        return export_census_block_groups()
    try:
        db: Session = next(get_db())
    except Exception as e:
        return jsonify({'error': f'Database connection failed: {str(e)}', 'data': []}), 500

    try:
        import requests
        from config.config import Config

        city = request.args.get('city')
        state = request.args.get('state')  # Used for geocoding + optional block-level filter
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)
        zip_code = request.args.get('zip_code')
        min_income = request.args.get('min_income', type=float)
        max_income = request.args.get('max_income', type=float)
        min_population = request.args.get('min_population', type=int)
        max_population = request.args.get('max_population', type=int)
        min_age = request.args.get('min_age', type=float)
        max_age = request.args.get('max_age', type=float)
        min_zoned_elementary_school_rating = request.args.get('min_zoned_elementary_school_rating', type=float)
        min_zoned_blended_school_rating = request.args.get('min_zoned_blended_school_rating', type=float)
        min_local_employment_score = request.args.get('min_local_employment_score', type=float)
        min_employment_access_score = request.args.get('min_employment_access_score', type=float)
        limit = request.args.get('limit', type=int, default=5000)

        # Resolve search to (lat, lng) or bbox for spatial query
        bbox = None  # (lng_min, lat_min, lng_max, lat_max)
        point = None  # (lng, lat)

        if lat is not None and lng is not None:
            point = (lng, lat)
        elif zip_code and str(zip_code).strip():
            # Geocode zip to point
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {'address': str(zip_code).strip(), 'key': Config.GOOGLE_MAPS_API_KEY, 'components': 'country:US'}
            r = requests.get(geocode_url, params=params, timeout=10)
            data = r.json()
            if data.get('status') == 'OK' and data.get('results'):
                loc = data['results'][0]['geometry']['location']
                point = (loc['lng'], loc['lat'])
        elif city and str(city).strip():
            # Geocode city to bbox (use viewport)
            addr = f"{city.strip()}, {state.strip()}, USA" if state else f"{city.strip()}, USA"
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {'address': addr, 'key': Config.GOOGLE_MAPS_API_KEY, 'components': 'country:US'}
            r = requests.get(geocode_url, params=params, timeout=10)
            data = r.json()
            if data.get('status') == 'OK' and data.get('results'):
                vp = data['results'][0].get('geometry', {}).get('viewport', {})
                ne = vp.get('northeast', {})
                sw = vp.get('southwest', {})
                if ne and sw:
                    bbox = (sw.get('lng'), sw.get('lat'), ne.get('lng'), ne.get('lat'))

        where_parts = ["geometry IS NOT NULL"]
        params: dict = {}
        if point:
            where_parts.append("ST_Contains(geometry, ST_SetSRID(ST_Point(:lng, :lat), 4326))")
            params["lng"] = point[0]
            params["lat"] = point[1]
        elif bbox:
            where_parts.append("ST_Intersects(geometry, ST_MakeEnvelope(:xmin, :ymin, :xmax, :ymax, 4326))")
            params["xmin"] = bbox[0]
            params["ymin"] = bbox[1]
            params["xmax"] = bbox[2]
            params["ymax"] = bbox[3]
        elif state and str(state).strip():
            # State-only: filter block groups by state (no city/zip/address)
            # census_block_groups.state uses FIPS codes (37=NC, 45=SC); accept both abbrev and FIPS
            where_parts.append("(UPPER(TRIM(state)) = UPPER(TRIM(:state_filter)) OR state = :state_fips)")
            params["state_filter"] = str(state).strip()
            _STATE_FIPS = {"NC": "37", "SC": "45", "VA": "51", "GA": "13", "TN": "47"}
            params["state_fips"] = _STATE_FIPS.get(str(state).strip().upper(), str(state).strip())
        else:
            return jsonify({"data": [], "total": 0, "limit": limit, "offset": 0})

        if min_income is not None:
            where_parts.append("average_household_income >= :min_income")
            params["min_income"] = min_income
        if max_income is not None:
            where_parts.append("average_household_income <= :max_income")
            params["max_income"] = max_income
        if min_population is not None:
            where_parts.append("population >= :min_population")
            params["min_population"] = min_population
        if max_population is not None:
            where_parts.append("population <= :max_population")
            params["max_population"] = max_population
        if min_age is not None:
            where_parts.append("median_age >= :min_age")
            params["min_age"] = min_age
        if max_age is not None:
            where_parts.append("median_age <= :max_age")
        if min_zoned_elementary_school_rating is not None:
            where_parts.append("se.rating IS NOT NULL AND se.rating >= :min_zoned_elementary_school_rating")
            params["min_zoned_elementary_school_rating"] = min_zoned_elementary_school_rating
        if min_zoned_blended_school_rating is not None:
            where_parts.append(
                "se.rating IS NOT NULL AND sm.rating IS NOT NULL AND sh.rating IS NOT NULL "
                "AND (se.rating + sm.rating + sh.rating) / 3.0 >= :min_zoned_blended_school_rating"
            )
            params["min_zoned_blended_school_rating"] = min_zoned_blended_school_rating
        if min_local_employment_score is not None:
            where_parts.append("cbg.local_employment_score IS NOT NULL AND cbg.local_employment_score >= :min_local_employment_score")
            params["min_local_employment_score"] = min_local_employment_score
        if min_employment_access_score is not None:
            where_parts.append("cbg.employment_access_score IS NOT NULL AND cbg.employment_access_score >= :min_employment_access_score")
            params["min_employment_access_score"] = min_employment_access_score
        # Filter block groups by state when provided with bbox/point (state-only adds it above)
        # census_block_groups.state uses FIPS codes (37=NC, 45=SC); accept both abbrev and FIPS
        if (point or bbox) and state and str(state).strip():
            where_parts.append("(UPPER(TRIM(state)) = UPPER(TRIM(:state_filter)) OR state = :state_fips)")
            params["state_filter"] = str(state).strip()
            _STATE_FIPS = {"NC": "37", "SC": "45", "VA": "51", "GA": "13", "TN": "47"}
            params["state_fips"] = _STATE_FIPS.get(str(state).strip().upper(), str(state).strip())

        where_sql = " AND ".join(where_parts)
        params["lim"] = limit
        qualified_where = where_sql
        for col in ("geometry", "state", "average_household_income", "population", "median_age"):
            qualified_where = re.sub(rf"\b{col}\b", f"cbg.{col}", qualified_where)
        # Employment score filters are already qualified with cbg., so ensure they're in qualified_where
        # (They're already there since they're in where_parts, but make sure)

        # Return geometry as GeoJSON; LEFT JOIN zoned schools and resolve names/ratings
        data_sql = text(f"""
            SELECT cbg.id, cbg.geoid, cbg.state, cbg.county, cbg.tract, cbg.block_group, cbg.population, cbg.median_age,
                   cbg.average_household_income, cbg.total_households, cbg.data_year,
                   cbg.local_employment_score, cbg.employment_access_score,
                   ST_AsGeoJSON(cbg.geometry)::json AS geometry,
                   z.zoned_elementary_school_id, z.zoned_middle_school_id, z.zoned_high_school_id,
                   se.name AS zoned_elementary_school_name, se.rating AS zoned_elementary_school_rating,
                   sm.name AS zoned_middle_school_name, sm.rating AS zoned_middle_school_rating,
                   sh.name AS zoned_high_school_name, sh.rating AS zoned_high_school_rating
            FROM census_block_groups cbg
            LEFT JOIN census_block_group_zoned_schools z ON z.geoid = cbg.geoid
            LEFT JOIN schools se ON se.id = z.zoned_elementary_school_id
            LEFT JOIN schools sm ON sm.id = z.zoned_middle_school_id
            LEFT JOIN schools sh ON sh.id = z.zoned_high_school_id
            WHERE {qualified_where}
            ORDER BY cbg.population DESC NULLS LAST
            LIMIT :lim
        """)
        rows = db.execute(data_sql, params).fetchall()

        # Count must use same JOINs when school filters are applied (se/sm/sh come from JOINs)
        # Use qualified_where when employment score filters are applied (they reference cbg.*)
        # Always use qualified_where when it differs from where_sql (has cbg. prefixes)
        if min_zoned_elementary_school_rating is not None or min_zoned_blended_school_rating is not None:
            # School filters require JOINs
            count_sql = text(f"""
                SELECT COUNT(*) FROM (
                    SELECT cbg.geoid
                    FROM census_block_groups cbg
                    LEFT JOIN census_block_group_zoned_schools z ON z.geoid = cbg.geoid
                    LEFT JOIN schools se ON se.id = z.zoned_elementary_school_id
                    LEFT JOIN schools sm ON sm.id = z.zoned_middle_school_id
                    LEFT JOIN schools sh ON sh.id = z.zoned_high_school_id
                    WHERE {qualified_where}
                ) sub
            """)
        else:
            # For employment filters or when qualified_where has cbg. prefixes, use qualified_where with cbg alias
            # This ensures employment filters (which use cbg.*) and state filters (qualified to cbg.state) work correctly
            count_sql = text(f"SELECT COUNT(*) FROM census_block_groups cbg WHERE {qualified_where}")
        total = db.execute(count_sql, params).scalar()

        # Use row._mapping for reliable key mapping (zip can misalign if SQL column order changes)
        data = []
        for row in rows:
            if hasattr(row, "_mapping"):
                d = dict(row._mapping)
                # Debug: log first row to see what keys are present
                if len(data) == 0:
                    print(f"[DEBUG] First row keys: {list(d.keys())}")
                    print(f"[DEBUG] LES value: {d.get('local_employment_score')} (type: {type(d.get('local_employment_score'))})")
                    print(f"[DEBUG] EAS value: {d.get('employment_access_score')} (type: {type(d.get('employment_access_score'))})")
                    # Check for alternative key names
                    for key in d.keys():
                        if 'employment' in key.lower() or 'score' in key.lower():
                            print(f"[DEBUG] Found key with 'employment' or 'score': {key} = {d[key]}")
            else:
                d = dict(zip(
                    ["id", "geoid", "state", "county", "tract", "block_group", "population", "median_age",
                     "average_household_income", "total_households", "data_year", 
                     "local_employment_score", "employment_access_score", "geometry",
                     "zoned_elementary_school_id", "zoned_middle_school_id", "zoned_high_school_id",
                     "zoned_elementary_school_name", "zoned_elementary_school_rating",
                     "zoned_middle_school_name", "zoned_middle_school_rating",
                     "zoned_high_school_name", "zoned_high_school_rating"], row))
            d["zip_code"] = None  # Block groups use geoid; frontend expects zip_code or geoid
            
            # Ensure scores are always present (even if None) - handle case where keys might be missing
            if "local_employment_score" not in d:
                d["local_employment_score"] = None
            if "employment_access_score" not in d:
                d["employment_access_score"] = None
            # Convert numeric scores to float for JSON serialization (handle Decimal, str, None)
            # PostgreSQL NUMERIC returns as Decimal, which needs explicit conversion
            from decimal import Decimal
            
            # Debug first row
            if len(data) == 0:
                print(f"[DEBUG] Before conversion - LES: {d.get('local_employment_score')} (type: {type(d.get('local_employment_score'))})")
                print(f"[DEBUG] Before conversion - EAS: {d.get('employment_access_score')} (type: {type(d.get('employment_access_score'))})")
            
            # Convert LES
            les_val = d.get("local_employment_score")
            if les_val is not None:
                try:
                    if isinstance(les_val, Decimal):
                        d["local_employment_score"] = float(les_val)
                    elif isinstance(les_val, (int, float)):
                        d["local_employment_score"] = float(les_val)
                    elif isinstance(les_val, str) and les_val.strip():
                        d["local_employment_score"] = float(les_val)
                    else:
                        d["local_employment_score"] = None
                except (ValueError, TypeError) as e:
                    print(f"[DEBUG] LES conversion error: {e}, value: {les_val}, type: {type(les_val)}")
                    d["local_employment_score"] = None
            else:
                d["local_employment_score"] = None
                
            # Convert EAS
            eas_val = d.get("employment_access_score")
            if eas_val is not None:
                try:
                    if isinstance(eas_val, Decimal):
                        d["employment_access_score"] = float(eas_val)
                    elif isinstance(eas_val, (int, float)):
                        d["employment_access_score"] = float(eas_val)
                    elif isinstance(eas_val, str) and eas_val.strip():
                        d["employment_access_score"] = float(eas_val)
                    else:
                        d["employment_access_score"] = None
                except (ValueError, TypeError) as e:
                    print(f"[DEBUG] EAS conversion error: {e}, value: {eas_val}, type: {type(eas_val)}")
                    d["employment_access_score"] = None
            else:
                d["employment_access_score"] = None
                
            if len(data) == 0:
                print(f"[DEBUG] After conversion - LES: {d.get('local_employment_score')} (type: {type(d.get('local_employment_score'))})")
                print(f"[DEBUG] After conversion - EAS: {d.get('employment_access_score')} (type: {type(d.get('employment_access_score'))})")
                
            data.append(d)

        return jsonify({"data": data, "total": total, "limit": limit, "offset": 0})
    except Exception as e:
        import traceback
        print(f"[ERROR] census-block-groups: {e}\n{traceback.format_exc()}")
        return jsonify({"error": str(e), "data": []}), 500


def _geojson_geom_to_shapefile_parts(geom: Dict[str, Any], max_vertices: int = 10000) -> List[List[Tuple[float, float]]]:
    """Convert GeoJSON geometry (Polygon or MultiPolygon) to list of rings for pyshp.
    Each ring is [(lng, lat), ...]. Simplifies if vertex count exceeds max_vertices (LandVision limit)."""
    try:
        from shapely.geometry import shape
        shp = shape(geom)
        if shp.is_empty:
            return []
        polys = list(shp.geoms) if hasattr(shp, 'geoms') else [shp]
        n = sum(len(list(p.exterior.coords)) for p in polys)
        if n > max_vertices and hasattr(shp, 'simplify'):
            shp = shp.simplify(0.00005, preserve_topology=True)
            polys = list(shp.geoms) if hasattr(shp, 'geoms') else [shp]
        parts = []
        for poly in polys:
            ext = list(poly.exterior.coords)
            if len(ext) < 3:
                continue
            if ext[0] != ext[-1]:
                ext.append(ext[0])
            parts.append([(float(x), float(y)) for x, y in ext])
        return parts
    except Exception:
        return []


@api.route('/export/test', methods=['GET'])
def export_test():
    """Simple test route to verify export endpoints are registered."""
    return jsonify({"ok": True, "message": "Export routes are working", "path": "/api/export/test"})


@api.route('/test-zip-download', methods=['GET'])
def api_test_zip_download():
    """Return a minimal valid ZIP. Use: http://127.0.0.1:5001/api/test-zip-download"""
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_STORED) as zf:
        zf.writestr('hello.txt', b'LandVision export test - if you see this, download is OK')
    zip_bytes = zip_buf.getvalue()
    from flask import Response
    r = Response(zip_bytes, mimetype='application/zip')
    r.headers['Content-Disposition'] = 'attachment; filename="test_download.zip"'
    r.headers['Content-Length'] = len(zip_bytes)
    return r


@api.route('/export/test-zip', methods=['GET'])
def export_test_zip():
    """Return a minimal valid ZIP so you can verify downloads are not corrupted."""
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_STORED) as zf:
        zf.writestr('hello.txt', b'LandVision export test - if you see this, download is OK')
    zip_bytes = zip_buf.getvalue()
    from flask import Response
    r = Response(zip_bytes, mimetype='application/zip')
    r.headers['Content-Disposition'] = 'attachment; filename="test_download.zip"'
    r.headers['Content-Length'] = len(zip_bytes)
    return r

@api.route('/census-block-groups/export', methods=['GET'])
@api.route('/export/landvision', methods=['GET'])
def export_census_block_groups():
    """
    Export current census block groups (same filters as map) for use in LandVision or other GIS.
    - format=geojson: GeoJSON FeatureCollection (same params as GET census-block-groups).
    - format=shapefile: ZIP containing .shp, .shx, .dbf, .prj (LandVision-ready: no nested folder, WGS84, max 30MB).
    """
    try:
        fmt = (request.args.get('format') or 'shapefile').strip().lower()
        if fmt not in ('geojson', 'shapefile'):
            return jsonify({"error": "format must be geojson or shapefile"}), 400

        # Call get_census_block_groups directly - it returns a Flask Response, so we need to extract JSON
        try:
            db: Session = next(get_db())
        except Exception as e:
            return jsonify({'error': f'Database connection failed: {str(e)}'}), 500

        try:
            import requests
            from config.config import Config

            # Build same query as get_census_block_groups
            city = request.args.get('city')
            state = request.args.get('state')
            lat = request.args.get('lat', type=float)
            lng = request.args.get('lng', type=float)
            zip_code = request.args.get('zip_code')
            min_income = request.args.get('min_income', type=float)
            max_income = request.args.get('max_income', type=float)
            min_population = request.args.get('min_population', type=int)
            max_population = request.args.get('max_population', type=int)
            min_age = request.args.get('min_age', type=float)
            max_age = request.args.get('max_age', type=float)
            min_zoned_elementary_school_rating = request.args.get('min_zoned_elementary_school_rating', type=float)
            min_zoned_blended_school_rating = request.args.get('min_zoned_blended_school_rating', type=float)
            min_local_employment_score = request.args.get('min_local_employment_score', type=float)
            min_employment_access_score = request.args.get('min_employment_access_score', type=float)
            limit = request.args.get('limit', type=int, default=10000)

            # Reuse the exact same logic from get_census_block_groups (geocoding, where building, query)
            bbox = None
            point = None

            if lat is not None and lng is not None:
                point = (lng, lat)
            elif zip_code and str(zip_code).strip():
                geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
                params_geo = {'address': str(zip_code).strip(), 'key': Config.GOOGLE_MAPS_API_KEY, 'components': 'country:US'}
                r = requests.get(geocode_url, params=params_geo, timeout=10)
                data_geo = r.json()
                if data_geo.get('status') == 'OK' and data_geo.get('results'):
                    loc = data_geo['results'][0]['geometry']['location']
                    point = (loc['lng'], loc['lat'])
            elif city and str(city).strip():
                addr = f"{city.strip()}, {state.strip()}, USA" if state else f"{city.strip()}, USA"
                geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
                params_geo = {'address': addr, 'key': Config.GOOGLE_MAPS_API_KEY, 'components': 'country:US'}
                r = requests.get(geocode_url, params=params_geo, timeout=10)
                data_geo = r.json()
                if data_geo.get('status') == 'OK' and data_geo.get('results'):
                    vp = data_geo['results'][0].get('geometry', {}).get('viewport', {})
                    ne = vp.get('northeast', {})
                    sw = vp.get('southwest', {})
                    if ne and sw:
                        bbox = (sw.get('lng'), sw.get('lat'), ne.get('lng'), ne.get('lat'))

            where_parts = ["geometry IS NOT NULL"]
            params: dict = {}
            if point:
                where_parts.append("ST_Contains(geometry, ST_SetSRID(ST_Point(:lng, :lat), 4326))")
                params["lng"] = point[0]
                params["lat"] = point[1]
            elif bbox:
                where_parts.append("ST_Intersects(geometry, ST_MakeEnvelope(:xmin, :ymin, :xmax, :ymax, 4326))")
                params["xmin"] = bbox[0]
                params["ymin"] = bbox[1]
                params["xmax"] = bbox[2]
                params["ymax"] = bbox[3]
            elif state and str(state).strip():
                where_parts.append("(UPPER(TRIM(state)) = UPPER(TRIM(:state_filter)) OR state = :state_fips)")
                params["state_filter"] = str(state).strip()
                _STATE_FIPS = {"NC": "37", "SC": "45", "VA": "51", "GA": "13", "TN": "47"}
                params["state_fips"] = _STATE_FIPS.get(str(state).strip().upper(), str(state).strip())
            else:
                db.close()
                return jsonify({"error": "No geographic search (city, zip, or address) provided"}), 400

            # Add filters (same as get_census_block_groups)
            if min_income is not None:
                where_parts.append("average_household_income >= :min_income")
                params["min_income"] = min_income
            if max_income is not None:
                where_parts.append("average_household_income <= :max_income")
                params["max_income"] = max_income
            if min_population is not None:
                where_parts.append("population >= :min_population")
                params["min_population"] = min_population
            if max_population is not None:
                where_parts.append("population <= :max_population")
                params["max_population"] = max_population
            if min_age is not None:
                where_parts.append("median_age >= :min_age")
                params["min_age"] = min_age
            if max_age is not None:
                where_parts.append("median_age <= :max_age")
                params["max_age"] = max_age
            if min_zoned_elementary_school_rating is not None:
                where_parts.append("se.rating IS NOT NULL AND se.rating >= :min_zoned_elementary_school_rating")
                params["min_zoned_elementary_school_rating"] = min_zoned_elementary_school_rating
            if min_zoned_blended_school_rating is not None:
                where_parts.append(
                    "se.rating IS NOT NULL AND sm.rating IS NOT NULL AND sh.rating IS NOT NULL "
                    "AND (se.rating + sm.rating + sh.rating) / 3.0 >= :min_zoned_blended_school_rating"
                )
                params["min_zoned_blended_school_rating"] = min_zoned_blended_school_rating
            if min_local_employment_score is not None:
                where_parts.append("cbg.local_employment_score IS NOT NULL AND cbg.local_employment_score >= :min_local_employment_score")
                params["min_local_employment_score"] = min_local_employment_score
            if min_employment_access_score is not None:
                where_parts.append("cbg.employment_access_score IS NOT NULL AND cbg.employment_access_score >= :min_employment_access_score")
                params["min_employment_access_score"] = min_employment_access_score
            if (point or bbox) and state and str(state).strip():
                where_parts.append("(UPPER(TRIM(state)) = UPPER(TRIM(:state_filter)) OR state = :state_fips)")
                params["state_filter"] = str(state).strip()
                _STATE_FIPS = {"NC": "37", "SC": "45", "VA": "51", "GA": "13", "TN": "47"}
                params["state_fips"] = _STATE_FIPS.get(str(state).strip().upper(), str(state).strip())

            where_sql = " AND ".join(where_parts)
            params["lim"] = limit
            qualified_where = where_sql
            for col in ("geometry", "state", "average_household_income", "population", "median_age"):
                qualified_where = re.sub(rf"\b{col}\b", f"cbg.{col}", qualified_where)

            # Execute query (same as get_census_block_groups)
            data_sql = text(f"""
                SELECT cbg.id, cbg.geoid, cbg.state, cbg.county, cbg.tract, cbg.block_group, cbg.population, cbg.median_age,
                       cbg.average_household_income, cbg.total_households, cbg.data_year,
                       cbg.local_employment_score, cbg.employment_access_score,
                       ST_AsGeoJSON(cbg.geometry)::json AS geometry,
                       z.zoned_elementary_school_id, z.zoned_middle_school_id, z.zoned_high_school_id,
                       se.name AS zoned_elementary_school_name, se.rating AS zoned_elementary_school_rating,
                       sm.name AS zoned_middle_school_name, sm.rating AS zoned_middle_school_rating,
                       sh.name AS zoned_high_school_name, sh.rating AS zoned_high_school_rating
                FROM census_block_groups cbg
                LEFT JOIN census_block_group_zoned_schools z ON z.geoid = cbg.geoid
                LEFT JOIN schools se ON se.id = z.zoned_elementary_school_id
                LEFT JOIN schools sm ON sm.id = z.zoned_middle_school_id
                LEFT JOIN schools sh ON sh.id = z.zoned_high_school_id
                WHERE {qualified_where}
                ORDER BY cbg.population DESC NULLS LAST
                LIMIT :lim
            """)
            rows = db.execute(data_sql, params).fetchall()

            # Convert rows to dicts (same as get_census_block_groups)
            data = []
            for row in rows:
                if hasattr(row, "_mapping"):
                    d = dict(row._mapping)
                else:
                    d = dict(zip(
                        ["id", "geoid", "state", "county", "tract", "block_group", "population", "median_age",
                         "average_household_income", "total_households", "data_year", 
                         "local_employment_score", "employment_access_score", "geometry",
                         "zoned_elementary_school_id", "zoned_middle_school_id", "zoned_high_school_id",
                         "zoned_elementary_school_name", "zoned_elementary_school_rating",
                         "zoned_middle_school_name", "zoned_middle_school_rating",
                         "zoned_high_school_name", "zoned_high_school_rating"], row))
                d["zip_code"] = None
                if "local_employment_score" not in d:
                    d["local_employment_score"] = None
                if "employment_access_score" not in d:
                    d["employment_access_score"] = None
                from decimal import Decimal
                les_val = d.get("local_employment_score")
                if les_val is not None:
                    try:
                        if isinstance(les_val, Decimal):
                            d["local_employment_score"] = float(les_val)
                        elif isinstance(les_val, (int, float)):
                            d["local_employment_score"] = float(les_val)
                        elif isinstance(les_val, str) and les_val.strip():
                            d["local_employment_score"] = float(les_val)
                        else:
                            d["local_employment_score"] = None
                    except (ValueError, TypeError):
                        d["local_employment_score"] = None
                else:
                    d["local_employment_score"] = None
                eas_val = d.get("employment_access_score")
                if eas_val is not None:
                    try:
                        if isinstance(eas_val, Decimal):
                            d["employment_access_score"] = float(eas_val)
                        elif isinstance(eas_val, (int, float)):
                            d["employment_access_score"] = float(eas_val)
                        elif isinstance(eas_val, str) and eas_val.strip():
                            d["employment_access_score"] = float(eas_val)
                        else:
                            d["employment_access_score"] = None
                    except (ValueError, TypeError):
                        d["employment_access_score"] = None
                else:
                    d["employment_access_score"] = None
                data.append(d)

            total = len(data)
            db.close()
        except Exception as e:
            try:
                db.close()
            except Exception:
                pass
            return jsonify({"error": f"Failed to fetch block groups: {str(e)}"}), 500

        if not data:
            return jsonify({
                "error": "No block groups to export. Search by city, zip, or address and apply filters first."
            }), 400

        if fmt == 'geojson':
            features = []
            for row in data:
                geom = row.get('geometry')
                if not geom:
                    continue
                props = {k: v for k, v in row.items() if k != 'geometry' and v is not None}
                features.append({"type": "Feature", "geometry": geom, "properties": props})
            fc = {"type": "FeatureCollection", "features": features}
            buf = io.BytesIO()
            import json
            buf.write(json.dumps(fc, separators=(',', ':')).encode('utf-8'))
            buf.seek(0)
            from datetime import datetime
            filename = f"block_groups_export_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.geojson"
            return send_file(
                buf,
                mimetype='application/geo+json',
                as_attachment=True,
                download_name=filename
            )

        # format=shapefile: build ZIP with .shp, .shx, .dbf, .prj (LandVision: no nested folder, WGS84)
        try:
            import shapefile
        except ImportError:
            return jsonify({"error": "shapefile export requires pyshp (pip install pyshp)"}), 500

        import tempfile
        import os
        tmpdir = tempfile.mkdtemp()
        base = os.path.join(tmpdir, 'block_groups')
        # pyshp 3.x: must pass target filepath (no extension)
        w = shapefile.Writer(target=base, shapeType=shapefile.POLYGON)
        w.autoBalance = 1
        # LandVision / DBF: 10-char field names
        w.field('geoid', 'C', size=12)
        w.field('tract', 'C', size=10)
        w.field('blk_grp', 'C', size=2)
        w.field('pop', 'N', size=8)
        w.field('med_age', 'N', size=6, decimal=1)
        w.field('mhi', 'N', size=10)
        w.field('elem_r', 'N', size=4, decimal=1)
        w.field('mid_r', 'N', size=4, decimal=1)
        w.field('high_r', 'N', size=4, decimal=1)
        w.field('les', 'N', size=4, decimal=2)
        w.field('eas', 'N', size=4, decimal=2)

        poly_count = 0
        for row in data:
            geom = row.get('geometry')
            if not geom:
                continue
            parts = _geojson_geom_to_shapefile_parts(geom)
            if not parts:
                continue
            w.poly(parts)
            re_ = row.get('zoned_elementary_school_rating')
            rm_ = row.get('zoned_middle_school_rating')
            rh_ = row.get('zoned_high_school_rating')
            les_ = row.get('local_employment_score')
            eas_ = row.get('employment_access_score')
            # Use 0 for missing numerics so DBF has no nulls (LandVision shape loader can fail on null)
            w.record(
                (row.get('geoid') or '')[:12],
                (row.get('tract') or '')[:10],
                (row.get('block_group') or '')[:2],
                int(row.get('population') or 0),
                round(float(row.get('median_age') or 0), 1),
                int(row.get('average_household_income') or 0),
                round(float(re_), 1) if re_ is not None else 0,
                round(float(rm_), 1) if rm_ is not None else 0,
                round(float(rh_), 1) if rh_ is not None else 0,
                round(float(les_), 2) if les_ is not None else 0,
                round(float(eas_), 2) if eas_ is not None else 0,
            )
            poly_count += 1

        if poly_count == 0:
            try:
                for ext in ('.shp', '.shx', '.dbf', '.prj'):
                    p = base + ext
                    if os.path.isfile(p):
                        os.remove(p)
            except Exception:
                pass
            try:
                os.rmdir(tmpdir)
            except Exception:
                pass
            return jsonify({
                "error": "No block group geometries could be converted to shapefile. Try a different search."
            }), 400

        w.close()
        try:
            # Add .prj for WGS84 (LandVision expects it)
            prj_path = base + '.prj'
            wkt = 'GEOGCS["WGS 84",DATUM["WGS_1984",SPHEROID["WGS 84",6378137,298.257223563]],PRIMEM["Greenwich",0],UNIT["degree",0.0174532925199433]]'
            with open(prj_path, 'w') as f:
                f.write(wkt)
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_STORED) as zf:
                for ext in ('.shp', '.shx', '.dbf', '.prj'):
                    path = base + ext
                    if os.path.isfile(path):
                        zf.write(path, os.path.basename(path))
            zip_bytes = zip_buf.getvalue()
            if len(zip_bytes) < 100:
                return jsonify({"error": "Generated ZIP is too small; shapefile may be invalid."}), 500
            from datetime import datetime
            filename = f"block_groups_landvision_{datetime.utcnow().strftime('%Y%m%d_%H%M')}.zip"
            from flask import Response
            resp = Response(zip_bytes, mimetype='application/zip')
            resp.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            resp.headers['Content-Length'] = len(zip_bytes)
            return resp
        finally:
            try:
                for ext in ('.shp', '.shx', '.dbf', '.prj'):
                    p = base + ext
                    if os.path.isfile(p):
                        os.remove(p)
            except Exception:
                pass
            try:
                os.rmdir(tmpdir)
            except Exception:
                pass

    except Exception as e:
        import traceback
        print(f"[ERROR] census-block-groups/export: {e}\n{traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500


@api.route('/debug/zoned-schools', methods=['GET'])
def debug_zoned_schools():
    """Diagnostic: verify census_block_group_zoned_schools has data and joins work."""
    try:
        db: Session = next(get_db())
        row = db.execute(text("""
            SELECT cbg.geoid, z.zoned_elementary_school_id, se.name AS elem_name, se.rating AS elem_rating
            FROM census_block_groups cbg
            LEFT JOIN census_block_group_zoned_schools z ON z.geoid = cbg.geoid
            LEFT JOIN schools se ON se.id = z.zoned_elementary_school_id
            WHERE z.zoned_elementary_school_id IS NOT NULL
            LIMIT 1
        """)).fetchone()
        if row:
            return jsonify({
                "ok": True,
                "message": "Zoned schools data found",
                "sample": {"geoid": row[0], "elem_id": row[1], "elem_name": row[2], "elem_rating": float(row[3]) if row[3] else None}
            })
        count = db.execute(text("SELECT COUNT(*) FROM census_block_group_zoned_schools")).scalar()
        return jsonify({"ok": False, "message": "No joined rows; census_block_group_zoned_schools has " + str(count) + " rows"})
    except Exception as e:
        import traceback
        return jsonify({"ok": False, "error": str(e), "traceback": traceback.format_exc()}), 500


@api.route('/census-data/zip/<zip_code>', methods=['GET'])
def get_census_data_by_zip(zip_code: str):
    """Get census data for a specific zip code. Fetches from Census API if not in database."""
    db: Session = next(get_db())
    
    row = db.execute(text(f"SELECT {_CENSUS_SQL_COLS} FROM census_data WHERE zip_code = :zip LIMIT 1"), {"zip": zip_code}).fetchone()
    
    if not row:
        # Try to fetch from Census API
        print(f"[INFO] Zip code {zip_code} not in database, fetching from Census API...")
        try:
            client = CensusAPIClient()
            census_data_list = client.fetch_zip_code_data([zip_code])
            
            if census_data_list and len(census_data_list) > 0:
                data = census_data_list[0]
                # Store in database (only columns that exist on model)
                new_record = CensusData(**_census_kwargs(data))
                db.add(new_record)
                db.commit()
                print(f"[INFO] Successfully fetched and stored census data for zip {zip_code}")
                return jsonify(new_record.to_dict())
            else:
                print(f"[WARN] Census API returned no data for zip {zip_code}")
                return jsonify({'error': 'Zip code not found in Census API'}), 404
        except Exception as e:
            print(f"[ERROR] Failed to fetch census data for zip {zip_code}: {e}")
            return jsonify({'error': f'Failed to fetch census data: {str(e)}'}), 500
    
    keys = ["id", "zip_code", "county", "population", "median_age", "average_household_income", "local_employment_rating", "data_year", "created_at", "updated_at", "total_schools", "elementary_schools", "middle_schools", "high_schools", "average_school_rating", "average_elementary_school_rating", "average_middle_school_rating", "average_high_school_rating"]
    d = dict(zip(keys, row))
    if d.get("local_employment_rating") is not None:
        d["local_employment_rating"] = float(d["local_employment_rating"])
    if d.get("created_at"):
        d["created_at"] = d["created_at"].isoformat() if hasattr(d["created_at"], "isoformat") else str(d["created_at"])
    if d.get("updated_at"):
        d["updated_at"] = d["updated_at"].isoformat() if hasattr(d["updated_at"], "isoformat") else str(d["updated_at"]) if d["updated_at"] else None
    return jsonify(d)

@api.route('/geocode-zip/<zip_code>', methods=['GET'])
def geocode_zip(zip_code: str):
    """Backend geocoding endpoint for zip codes."""
    try:
        import requests
        from config.config import Config
        
        # Use Google Geocoding API via backend
        # This helps if frontend API key has restrictions
        geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
        params = {
            'address': zip_code,
            'key': Config.GOOGLE_MAPS_API_KEY,
            'components': 'country:US'
        }
        
        response = requests.get(geocode_url, params=params, timeout=10)
        data = response.json()
        
        if data['status'] == 'OK' and data['results']:
            result = data['results'][0]
            geometry = result['geometry']
            
            return jsonify({
                'success': True,
                'location': {
                    'lat': geometry['location']['lat'],
                    'lng': geometry['location']['lng']
                },
                'bounds': geometry.get('bounds'),
                'viewport': geometry.get('viewport')
            })
        else:
            return jsonify({
                'success': False,
                'error': data.get('error_message', data.get('status', 'Unknown error'))
            }), 400
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@api.route('/zip-boundary/<zip_code>', methods=['GET'])
def get_zip_boundary(zip_code: str):
    """Get GeoJSON boundary polygon for a zip code."""
    try:
        import requests
        import json
        from pathlib import Path
        
        # FIRST: Check if we have a locally stored boundary (fastest, most reliable)
        try:
            boundary_file = Path('data/zip_boundaries') / f"{zip_code}.geojson"
            if boundary_file.exists():
                with open(boundary_file, 'r') as f:
                    data = json.load(f)
                    return jsonify(data)
        except Exception as e:
            print(f"Local boundary file check failed: {e}")
        
        # Try multiple sources for zip code boundaries
        # Source 1: OpenDataSoft API (most reliable)
        try:
            ods_url = "https://public.opendatasoft.com/api/explore/v2.1/catalog/datasets/us-zip-code-labels-and-boundaries/records"
            # Try with ZCTA5CE10 field (Census format)
            params = {
                'where': f'zcta5ce10="{zip_code}"',
                'limit': 1,
                'select': 'zcta5ce10,geo_shape'
            }
            response = requests.get(ods_url, params=params, timeout=15)
            if response.status_code == 200:
                data = response.json()
                if 'results' in data and len(data['results']) > 0:
                    record = data['results'][0]
                    if 'record' in record and 'fields' in record['record']:
                        fields = record['record']['fields']
                        if 'geo_shape' in fields:
                            geometry = fields['geo_shape']
                            geojson = {
                                'type': 'FeatureCollection',
                                'features': [{
                                    'type': 'Feature',
                                    'geometry': geometry,
                                    'properties': {'ZCTA5CE10': zip_code}
                                }]
                            }
                            try:
                                boundary_dir = Path('data/zip_boundaries')
                                boundary_dir.mkdir(parents=True, exist_ok=True)
                                boundary_file = boundary_dir / f"{zip_code}.geojson"
                                with open(boundary_file, 'w') as f:
                                    json.dump(geojson, f)
                            except Exception:
                                pass
                            return jsonify(geojson)
            
            # Try with zip_code field
            params2 = {
                'where': f'zip_code="{zip_code}"',
                'limit': 1,
                'select': 'zip_code,geo_shape'
            }
            response2 = requests.get(ods_url, params=params2, timeout=15)
            if response2.status_code == 200:
                data2 = response2.json()
                if 'results' in data2 and len(data2['results']) > 0:
                    record = data2['results'][0]
                    if 'record' in record and 'fields' in record['record']:
                        fields = record['record']['fields']
                        if 'geo_shape' in fields:
                            geometry = fields['geo_shape']
                            geojson = {
                                'type': 'FeatureCollection',
                                'features': [{
                                    'type': 'Feature',
                                    'geometry': geometry,
                                    'properties': {'ZCTA5CE10': zip_code}
                                }]
                            }
                            try:
                                boundary_dir = Path('data/zip_boundaries')
                                boundary_dir.mkdir(parents=True, exist_ok=True)
                                boundary_file = boundary_dir / f"{zip_code}.geojson"
                                with open(boundary_file, 'w') as f:
                                    json.dump(geojson, f)
                            except Exception:
                                pass
                            return jsonify(geojson)
        except Exception as e:
            print(f"OpenDataSoft failed: {e}")
        
        # Source 2: Try boundaries.io API (FREE tier available)
        try:
            from config.config import Config
            boundaries_api_key = getattr(Config, 'BOUNDARIES_IO_API_KEY', None)
            
            if boundaries_api_key:
                # Use Boundaries.io API (requires free API key)
                boundaries_url = f"https://boundaries.io/api/v1/boundary"
                params = {
                    'zipcode': zip_code,
                    'api_key': boundaries_api_key,
                    'format': 'geojson'
                }
                response = requests.get(boundaries_url, params=params, timeout=15)
                if response.status_code == 200:
                    data = response.json()
                    # Save locally for future use
                    try:
                        boundary_dir = Path('data/zip_boundaries')
                        boundary_dir.mkdir(parents=True, exist_ok=True)
                        boundary_file = boundary_dir / f"{zip_code}.geojson"
                        with open(boundary_file, 'w') as f:
                            json.dump(data, f)
                    except:
                        pass
                    return jsonify(data)
            else:
                # Try free endpoint (may have rate limits)
                boundaries_url = f"https://boundaries-io.herokuapp.com/zip/{zip_code}"
                response = requests.get(boundaries_url, timeout=15)
                if response.status_code == 200:
                    data = response.json()
                    if 'boundaries' in data or 'geometry' in data:
                        return jsonify(data)
        except Exception as e:
            print(f"Boundaries.io failed: {e}")
        
        # Source 3: Try GitHub repository (state-based)
        try:
            # Try different GitHub sources
            github_sources = [
                f"https://raw.githubusercontent.com/OpenDataDE/State-zip-code-GeoJSON/master/{zip_code[0]}/{zip_code}_polygon.geojson",
                f"https://raw.githubusercontent.com/OpenDataDE/State-zip-code-GeoJSON/master/zcta5/{zip_code}_polygon.geojson",
            ]
            
            for github_url in github_sources:
                response = requests.get(github_url, timeout=10)
                if response.status_code == 200:
                    data = response.json()
                    if 'type' in data and (data['type'] == 'FeatureCollection' or data['type'] == 'Feature'):
                        if data['type'] == 'Feature':
                            # Wrap single feature in FeatureCollection
                            return jsonify({
                                'type': 'FeatureCollection',
                                'features': [data]
                            })
                        return jsonify(data)
        except Exception as e:
            print(f"GitHub source failed: {e}")
        
        # Source 4: Try using Census TIGERweb (ZCTA layer is 2, not 82!)
        try:
            feature_url = "https://tigerweb.geo.census.gov/arcgis/rest/services/TIGERweb/tigerWMS_Current/MapServer/2/query"
            
            # Use objectIds approach - first find the object ID
            # Then query by object ID
            params_list = [
                {'where': f"ZCTA5='{zip_code}'", 'outFields': '*', 'f': 'geojson', 'outSR': '4326', 'returnGeometry': 'true'},
                {'where': f"ZCTA5 = '{zip_code}'", 'outFields': '*', 'f': 'geojson', 'outSR': '4326', 'returnGeometry': 'true'},
                {'where': f"ZCTA5CE10='{zip_code}'", 'outFields': '*', 'f': 'geojson', 'outSR': '4326', 'returnGeometry': 'true'},
            ]
            
            for params in params_list:
                try:
                    response = requests.get(feature_url, params=params, timeout=25)
                    if response.status_code == 200:
                        data = response.json()
                        # Check for actual features, not errors
                        if 'features' in data and isinstance(data['features'], list) and len(data['features']) > 0:
                            # Validate it's actually a polygon
                            feature = data['features'][0]
                            if 'geometry' in feature and feature['geometry'].get('type') in ['Polygon', 'MultiPolygon']:
                                # Save locally for future use
                                try:
                                    boundary_dir = Path('data/zip_boundaries')
                                    boundary_dir.mkdir(parents=True, exist_ok=True)
                                    boundary_file = boundary_dir / f"{zip_code}.geojson"
                                    with open(boundary_file, 'w') as f:
                                        json.dump(data, f)
                                except:
                                    pass
                                return jsonify(data)
                except:
                    continue
        except Exception as e:
            print(f"Census TIGER failed: {e}")
        
        # Source 5: Try alternative Census TIGERweb endpoints (all FREE)
        census_endpoints = [
            "https://tigerweb.geo.census.gov/arcgis/rest/services/TIGERweb/tigerWMS_ACS2022/MapServer/2/query",
            "https://tigerweb.geo.census.gov/arcgis/rest/services/TIGERweb/tigerWMS_ACS2021/MapServer/2/query",
        ]
        
        for alt_url in census_endpoints:
            try:
                params = {
                    'where': f"ZCTA5='{zip_code}'",
                    'outFields': '*',
                    'f': 'geojson',
                    'outSR': '4326',
                    'returnGeometry': 'true'
                }
                response = requests.get(alt_url, params=params, timeout=20)
                if response.status_code == 200:
                    data = response.json()
                    if 'features' in data and len(data['features']) > 0:
                        try:
                            boundary_dir = Path('data/zip_boundaries')
                            boundary_dir.mkdir(parents=True, exist_ok=True)
                            boundary_file = boundary_dir / f"{zip_code}.geojson"
                            with open(boundary_file, 'w') as f:
                                json.dump(data, f)
                        except Exception:
                            pass
                        return jsonify(data)
            except Exception as e:
                print(f"Census endpoint {alt_url} failed: {e}")
                continue
        
        # If all sources fail, return 404 (frontend will use approximate boundary)
        return jsonify({
            'error': 'Boundary not found',
            'message': f'Could not fetch exact boundary for zip code {zip_code}. Using approximate boundary from geocoding.'
        }), 404
        
    except Exception as e:
        print(f"Error in get_zip_boundary: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@api.route('/census-data', methods=['POST'])
def add_census_data():
    """Add or update census data."""
    db: Session = next(get_db())
    
    data = request.get_json()
    
    if not data or 'zip_code' not in data:
        return jsonify({'error': 'zip_code is required'}), 400
    
    # Check if record exists (load_only so we never SELECT city - column may not exist in DB yet)
    existing = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(
        CensusData.zip_code == data['zip_code']
    ).first()
    
    if existing:
        # Update existing record
        for key, value in data.items():
            if hasattr(existing, key) and key != 'id':
                setattr(existing, key, value)
    else:
        # Create new record (only columns that exist on model)
        existing = CensusData(**_census_kwargs(data))
        db.add(existing)
    
    db.commit()
    db.refresh(existing)
    
    return jsonify(existing.to_dict()), 201 if not existing.id else 200

@api.route('/census-data/bulk', methods=['POST'])
def add_census_data_bulk():
    """Add multiple census records at once."""
    db: Session = next(get_db())
    
    data_list = request.get_json()
    
    if not isinstance(data_list, list):
        return jsonify({'error': 'Expected a list of records'}), 400
    
    added = 0
    updated = 0
    
    for data in data_list:
        if 'zip_code' not in data:
            continue
        
        existing = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(
            CensusData.zip_code == data['zip_code']
        ).first()
        
        if existing:
            for key, value in data.items():
                if hasattr(existing, key) and key != 'id':
                    setattr(existing, key, value)
            updated += 1
        else:
            new_record = CensusData(**_census_kwargs(data))
            db.add(new_record)
            added += 1
    
    db.commit()
    
    return jsonify({
        'message': 'Bulk update completed',
        'added': added,
        'updated': updated
    })

@api.route('/census-data/fetch', methods=['POST'])
def fetch_census_data():
    """Fetch census data from Census Bureau API and store in database."""
    db: Session = next(get_db())
    
    request_data = request.get_json() or {}
    zip_codes = request_data.get('zip_codes')  # Optional list of zip codes
    
    client = CensusAPIClient()
    census_data = client.fetch_zip_code_data(zip_codes)
    
    if not census_data:
        return jsonify({'error': 'No data fetched from Census API'}), 400
    
    # Store in database
    added = 0
    updated = 0
    
    for data in census_data:
        existing = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS)).filter(
            CensusData.zip_code == data['zip_code']
        ).first()
        
        if existing:
            for key, value in data.items():
                if hasattr(existing, key):
                    setattr(existing, key, value)
            updated += 1
        else:
            new_record = CensusData(**_census_kwargs(data))
            db.add(new_record)
            added += 1
    
    db.commit()
    
    return jsonify({
        'message': 'Census data fetched and stored',
        'added': added,
        'updated': updated,
        'total_fetched': len(census_data)
    })

@api.route('/export/csv', methods=['GET'])
def export_to_csv():
    """Export census data to CSV file (download)."""
    from flask import Response
    import csv
    import io
    from datetime import datetime
    
    try:
        db: Session = next(get_db())
        
        # Get filter parameters from request (same as /api/census-data endpoint)
        zip_code = request.args.get('zip_code')
        min_income = request.args.get('min_income', type=float)
        max_income = request.args.get('max_income', type=float)
        min_population = request.args.get('min_population', type=int)
        max_population = request.args.get('max_population', type=int)
        limit = request.args.get('limit', type=int, default=10000)
        
        # Build query with filters (load_only so we never SELECT city - column may not exist in DB yet)
        query = db.query(CensusData).options(load_only(*_CENSUS_LOAD_COLUMNS))
        
        if zip_code:
            query = query.filter(CensusData.zip_code == zip_code)
        if min_income:
            query = query.filter(CensusData.average_household_income >= min_income)
        if max_income:
            query = query.filter(CensusData.average_household_income <= max_income)
        if min_population:
            query = query.filter(CensusData.population >= min_population)
        if max_population:
            query = query.filter(CensusData.population <= max_population)
        
        # Apply limit and get records
        records = query.limit(limit).all()
        
        # Convert to list of dicts
        data = [record.to_dict() for record in records]
        
        if not data:
            # Provide more helpful error message
            if zip_code:
                return jsonify({
                    'error': 'No data to export',
                    'message': f'Zip code {zip_code} not found in database. This zip code may not have census data available.',
                    'zip_code': zip_code
                }), 400
            else:
                return jsonify({
                    'error': 'No data to export',
                    'message': 'No records match the specified filters.'
                }), 400
        
        # Create CSV in memory
        output = io.StringIO()
        
        # Get headers from first record
        headers = list(data[0].keys())
        writer = csv.DictWriter(output, fieldnames=headers)
        
        # Write headers
        writer.writeheader()
        
        # Write data rows
        for record in data:
            writer.writerow(record)
        
        # Create filename with timestamp and zip code if applicable
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if zip_code:
            filename = f'census_data_{zip_code}_{timestamp}.csv'
        else:
            filename = f'census_data_{timestamp}.csv'
        
        # Create response with CSV data
        response = Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"'
            }
        )
        
        return response
        
    except ImportError as ie:
        missing_lib = str(ie)
        if 'googleapiclient' in missing_lib or 'google-api-python-client' in missing_lib.lower():
            return jsonify({
                'error': 'Google API client library not installed',
                'message': 'Install with: pip install google-api-python-client'
            }), 500
        else:
            return jsonify({
                'error': 'Google Sheets libraries not installed',
                'message': 'Install with: pip install gspread google-auth google-auth-oauthlib google-api-python-client'
            }), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@api.route('/export/report', methods=['GET'])
def export_report():
    """Export site report as Word document or PDF with demographics and top 10 schools."""
    try:
        from flask import Response
        from datetime import datetime
        from io import BytesIO
        import requests
        from config.config import Config
        from sqlalchemy import text
        
        # Get parameters
        address = request.args.get('address')
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)
        zip_code = request.args.get('zip_code')
        format_type = request.args.get('format', 'docx')  # 'docx' or 'pdf'
        
        if not address:
            return jsonify({'error': 'Address parameter is required'}), 400
        
        # If lat/lng not provided, geocode the address
        if lat is None or lng is None:
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {
                'address': address,
                'key': Config.GOOGLE_MAPS_API_KEY,
                'components': 'country:US'
            }
            response = requests.get(geocode_url, params=params, timeout=10)
            data = response.json()
            
            if data['status'] != 'OK' or not data['results']:
                return jsonify({
                    'error': 'Could not geocode address',
                    'details': data.get('error_message', data.get('status', 'Unknown error'))
                }), 400
            
            location = data['results'][0]['geometry']['location']
            lat = location['lat']
            lng = location['lng']
            
            # Extract zip code if not provided
            if not zip_code:
                for component in data['results'][0]['address_components']:
                    if 'postal_code' in component['types']:
                        zip_code = component['long_name']
                        break
        
        db: Session = next(get_db())
        
        # Get census data for zip code
        census_record = None
        if zip_code:
            row = db.execute(text(f"SELECT {_CENSUS_SQL_COLS} FROM census_data WHERE zip_code = :zip LIMIT 1"), {"zip": zip_code}).fetchone()
            if row:
                keys = ["id", "zip_code", "county", "population", "median_age", "average_household_income", "local_employment_rating", "data_year", "created_at", "updated_at", "total_schools", "elementary_schools", "middle_schools", "high_schools", "average_school_rating", "average_elementary_school_rating", "average_middle_school_rating", "average_high_school_rating"]
                census_record = dict(zip(keys, row))
                if census_record.get("local_employment_rating") is not None:
                    census_record["local_employment_rating"] = float(census_record["local_employment_rating"])
        
        # STEP 1: Get "zoned" schools = nearest elementary, middle, high in school_data within ~5 miles (no Apify)
        zoned_schools = []
        zoned_school_names = set()
        search_radius_zoned = 5.0 / 69.0
        qp = {
            'lat': lat, 'lng': lng,
            'lat_min': lat - search_radius_zoned, 'lat_max': lat + search_radius_zoned,
            'lng_min': lng - search_radius_zoned, 'lng_max': lng + search_radius_zoned
        }
        dist_sql = "3959 * acos(cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) + sin(radians(:lat)) * sin(radians(latitude)))"
        q_elem = text("SELECT elementary_school_name, elementary_school_address, elementary_school_rating, latitude, longitude FROM school_data WHERE elementary_school_rating IS NOT NULL AND latitude BETWEEN :lat_min AND :lat_max AND longitude BETWEEN :lng_min AND :lng_max ORDER BY " + dist_sql + " LIMIT 1")
        q_mid = text("SELECT middle_school_name, middle_school_address, middle_school_rating, latitude, longitude FROM school_data WHERE middle_school_rating IS NOT NULL AND latitude BETWEEN :lat_min AND :lat_max AND longitude BETWEEN :lng_min AND :lng_max ORDER BY " + dist_sql + " LIMIT 1")
        q_high = text("SELECT high_school_name, high_school_address, high_school_rating, latitude, longitude FROM school_data WHERE high_school_rating IS NOT NULL AND latitude BETWEEN :lat_min AND :lat_max AND longitude BETWEEN :lng_min AND :lng_max ORDER BY " + dist_sql + " LIMIT 1")
        import math
        for query, name_col, addr_col, rating_col, school_type in [
            (q_elem, 0, 1, 2, 'Elementary'),
            (q_mid, 0, 1, 2, 'Middle'),
            (q_high, 0, 1, 2, 'High'),
        ]:
            row = db.execute(query, qp).fetchone()
            if row and row[name_col]:
                sch_lat, sch_lng = row[3], row[4]
                dist = 3959 * math.acos(
                    math.cos(math.radians(lat)) * math.cos(math.radians(sch_lat)) * math.cos(math.radians(sch_lng) - math.radians(lng)) +
                    math.sin(math.radians(lat)) * math.sin(math.radians(sch_lat))
                )
                zoned_schools.append({
                    'name': row[name_col],
                    'address': row[addr_col] or 'N/A',
                    'type': school_type,
                    'rating': row[rating_col],
                    'distance': dist,
                    'is_zoned': True
                })
                zoned_school_names.add(row[name_col].lower())

        # STEP 2: Get additional schools (sorted by rating, then distance) to fill up to 10 total
        # Mirror GreatSchools logic: 5-7 miles radius, public/charter only, rated schools only
        search_radius = 6.0 / 69.0  # ~6 miles radius (GreatSchools typically uses 5-7 miles)
        additional_query = text("""
            SELECT 
                elementary_school_name, elementary_school_address, elementary_school_rating,
                middle_school_name, middle_school_address, middle_school_rating,
                high_school_name, high_school_address, high_school_rating,
                latitude, longitude,
                3959 * acos(
                    cos(radians(:lat)) * cos(radians(latitude)) * 
                    cos(radians(longitude) - radians(:lng)) + 
                    sin(radians(:lat)) * sin(radians(latitude))
                ) as distance_miles
            FROM school_data
            WHERE (elementary_school_rating IS NOT NULL 
                   OR middle_school_rating IS NOT NULL 
                   OR high_school_rating IS NOT NULL)
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
        """)
        
        query_params = {
            'lat': lat,
            'lng': lng,
            'lat_min': lat - search_radius,
            'lat_max': lat + search_radius,
            'lng_min': lng - search_radius,
            'lng_max': lng + search_radius
        }
        
        all_results = db.execute(additional_query, query_params).fetchall()
        
        # Extract all schools, excluding zoned ones
        additional_schools = []
        for row in all_results:
            distance = row[11]
            
            # Elementary
            if row[0] and row[2] is not None and row[0].lower() not in zoned_school_names:
                additional_schools.append({
                    'name': row[0],
                    'address': row[1] or 'N/A',
                    'type': 'Elementary',
                    'rating': row[2],
                    'distance': distance,
                    'is_zoned': False
                })
            
            # Middle
            if row[3] and row[5] is not None and row[3].lower() not in zoned_school_names:
                additional_schools.append({
                    'name': row[3],
                    'address': row[4] or 'N/A',
                    'type': 'Middle',
                    'rating': row[5],
                    'distance': distance,
                    'is_zoned': False
                })
            
            # High
            if row[6] and row[8] is not None and row[6].lower() not in zoned_school_names:
                additional_schools.append({
                    'name': row[6],
                    'address': row[7] or 'N/A',
                    'type': 'High',
                    'rating': row[8],
                    'distance': distance,
                    'is_zoned': False
                })
        
        # Sort additional schools by rating (descending), then distance (ascending)
        # Mirror GreatSchools: Filter for public/charter schools only, exclude private and unrated
        # Note: We'll filter by school type if we have that data, but for now we'll rely on rating filter
        additional_schools = [s for s in additional_schools if s['rating'] is not None]  # Exclude unrated schools
        additional_schools.sort(key=lambda x: (-x['rating'] if x['rating'] is not None else 0, x['distance'] if x['distance'] is not None else float('inf')))
        
        # Combine: zoned schools first, then additional schools up to 10 total
        schools = zoned_schools + additional_schools[:10 - len(zoned_schools)]
        
        # Format for display
        formatted_schools = []
        for school in schools:
            formatted_schools.append({
                'name': school['name'],
                'address': school['address'],
                'type': school['type'],
                'rating': f"{school['rating']:.1f}/10" if school['rating'] is not None else 'N/A',
                'distance': f"{school['distance']:.2f} miles" if school['distance'] is not None else 'N/A',
                'is_zoned': school.get('is_zoned', False)
            })
        
        schools = formatted_schools
        
        # Generate document
        if format_type == 'pdf':
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a73e8'),
                spaceAfter=30
            )
            story.append(Paragraph("Site Selection Report", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Demographics Section
            story.append(Paragraph("<b>1. Demographics</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            demo_data = [
                ['Field', 'Value'],
                ['Address', address],
                ['Zip Code', zip_code or 'N/A'],
                ['Population', f"{census_record['population']:,}" if census_record and census_record.get('population') else 'N/A'],
                ['Median Household Income (MHI)', f"${census_record['average_household_income']:,.0f}" if census_record and census_record.get('average_household_income') else 'N/A'],
                ['Median Age', f"{census_record['median_age']:.1f} years" if census_record and census_record.get('median_age') else 'N/A'],
                ['Local Employment Rating', f"{census_record['local_employment_rating']:.1f} / 10" if census_record and census_record.get('local_employment_rating') is not None else 'N/A']
            ]
            
            demo_table = Table(demo_data, colWidths=[2.5*inch, 4*inch])
            demo_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(demo_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Schools Section
            story.append(Paragraph("<b>2. Great School Scores (Top 10 Schools)</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            if schools:
                # Create table with styled cells for zoned schools
                from reportlab.platypus import Table, TableStyle, Paragraph
                from reportlab.lib.styles import ParagraphStyle
                
                school_data = [['School Name', 'Address', 'Type', 'Rating', 'Proximity']]
                zoned_row_indices = []  # Track which rows are zoned
                
                for i, school in enumerate(schools):
                    row_idx = i + 1  # +1 for header row
                    school_data.append([
                        school['name'],
                        school['address'],
                        school['type'],
                        school['rating'],
                        school['distance']
                    ])
                    if school.get('is_zoned', False):
                        zoned_row_indices.append(row_idx)
                
                school_table = Table(school_data, colWidths=[1.8*inch, 1.8*inch, 0.8*inch, 0.8*inch, 0.8*inch])
                
                # Base table style
                table_style = [
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]
                
                # Style zoned schools with bold and blue color
                for row_idx in zoned_row_indices:
                    table_style.extend([
                        ('FONTNAME', (0, row_idx), (-1, row_idx), 'Helvetica-Bold'),
                        ('TEXTCOLOR', (0, row_idx), (-1, row_idx), colors.HexColor('#1a73e8')),
                        ('BACKGROUND', (0, row_idx), (-1, row_idx), colors.HexColor('#e8f0fe')),
                    ])
                
                # Add alternating row colors for non-zoned rows
                non_zoned_rows = [i for i in range(1, len(school_data)) if i not in zoned_row_indices]
                for i, row_idx in enumerate(non_zoned_rows):
                    bg_color = colors.white if i % 2 == 0 else colors.lightgrey
                    table_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), bg_color))
                
                school_table.setStyle(TableStyle(table_style))
                story.append(school_table)
            else:
                story.append(Paragraph("No school data available for this location.", styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'site_report_{zip_code or "unknown"}_{timestamp}.pdf'
            
            return Response(
                buffer.getvalue(),
                mimetype='application/pdf',
                headers={'Content-Disposition': f'attachment; filename="{filename}"'}
            )
        else:
            # Word document (docx)
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Title
                title = doc.add_heading('Site Selection Report', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.runs[0].font.color.rgb = RGBColor(26, 115, 232)
                
                # Demographics Section
                doc.add_heading('1. Demographics', 1)
                
                demo_data = [
                    ('Field', 'Value'),
                    ('Address', address),
                    ('Zip Code', zip_code or 'N/A'),
                    ('Population', f"{census_record['population']:,}" if census_record and census_record.get('population') else 'N/A'),
                    ('Median Household Income (MHI)', f"${census_record['average_household_income']:,.0f}" if census_record and census_record.get('average_household_income') else 'N/A'),
                    ('Median Age', f"{census_record['median_age']:.1f} years" if census_record and census_record.get('median_age') else 'N/A'),
                    ('Local Employment Rating', f"{census_record['local_employment_rating']:.1f} / 10" if census_record and census_record.get('local_employment_rating') is not None else 'N/A')
                ]
                
                demo_table = doc.add_table(rows=6, cols=2)
                demo_table.style = 'Light Grid Accent 1'
                
                for i, (field, value) in enumerate(demo_data):
                    row = demo_table.rows[i]
                    row.cells[0].text = field
                    row.cells[1].text = str(value)
                    if i == 0:  # Header row
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                
                doc.add_paragraph()  # Spacing
                
                # Schools Section
                doc.add_heading('2. Great School Scores (Top 10 Schools)', 1)
                
                if schools:
                    school_table = doc.add_table(rows=1, cols=5)
                    school_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = school_table.rows[0].cells
                    headers = ['School Name', 'Address', 'Type', 'Rating', 'Proximity']
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data rows - style zoned schools with bold and blue color
                    for school in schools:
                        row_cells = school_table.add_row().cells
                        is_zoned = school.get('is_zoned', False)
                        
                        # Set cell values and style
                        for i, value in enumerate([school['name'], school['address'], school['type'], school['rating'], school['distance']]):
                            row_cells[i].text = value
                            
                            # Style zoned schools
                            if is_zoned:
                                for paragraph in row_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(26, 115, 232)  # Blue color
                                    # Set cell background color (light blue)
                                    shading_elm = paragraph._element.get_or_add_pPr().get_or_add_shd()
                                    shading_elm.set('fill', 'E8F0FE')  # Light blue background
                else:
                    doc.add_paragraph('No school data available for this location.')
                
                # Save to BytesIO
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'site_report_{zip_code or "unknown"}_{timestamp}.docx'
                
                return Response(
                    buffer.getvalue(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': f'attachment; filename="{filename}"'}
                )
            except ImportError:
                return jsonify({
                    'error': 'python-docx not installed',
                    'message': 'Install with: pip install python-docx'
                }), 500
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in export_report: {e}")
        print(error_trace)
        return jsonify({
            'error': str(e),
            'message': 'Error generating report',
            'traceback': error_trace
        }), 500


@api.route('/schools/address', methods=['GET'])
def get_schools_by_address():
    """Get school ratings for an address. Prefers zoned schools (attendance boundaries) for NC/SC, else nearest schools in school_data."""
    try:
        import requests
        from config.config import Config
        from sqlalchemy import text

        address = request.args.get('address')
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)

        if not address:
            return jsonify({'error': 'Address parameter is required'}), 400

        # Geocode if lat/lng not provided
        if lat is None or lng is None:
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {
                'address': address,
                'key': Config.GOOGLE_MAPS_API_KEY,
                'components': 'country:US'
            }
            response = requests.get(geocode_url, params=params, timeout=10)
            data = response.json()
            if data['status'] != 'OK' or not data['results']:
                return jsonify({
                    'error': 'Could not geocode address',
                    'details': data.get('error_message', data.get('status', 'Unknown error'))
                }), 400
            location = data['results'][0]['geometry']['location']
            lat = location['lat']
            lng = location['lng']

        zip_code = None
        if address:
            import re
            zip_match = re.search(r'\b\d{5}(?:-\d{4})?\b', address)
            if zip_match:
                zip_code = zip_match.group(0)

        db: Session = next(get_db())

        # STEP 1: Try zoned schools (attendance boundaries) for NC/SC
        zone_ids_for_zip = None
        if zip_code:
            try:
                rows = db.execute(
                    text("SELECT zone_id FROM zone_zips WHERE zip_code = :zip"),
                    {"zip": zip_code}
                ).fetchall()
                zone_ids_for_zip = [r[0] for r in rows] if rows else []
            except Exception:
                zone_ids_for_zip = None

        if zone_ids_for_zip is not None and len(zone_ids_for_zip) > 0:
            zones = db.query(AttendanceZone).filter(
                AttendanceZone.id.in_(zone_ids_for_zip)
            ).all()
        else:
            zones = db.query(AttendanceZone).filter(
                or_(AttendanceZone.state == 'NC', AttendanceZone.state == 'SC')
            ).all()

        if zones:
            zones_list = [z.to_dict() for z in zones]
            by_level = find_all_zoned_schools(lat, lng, zones_list)
            elem_zones = by_level.get('elementary', [])
            mid_zones = by_level.get('middle', [])
            high_zones = by_level.get('high', [])

            if elem_zones or mid_zones or high_zones:
                def first_zoned(zone_list, level_key):
                    if not zone_list:
                        return None, None, None
                    z = zone_list[0]
                    name = z.get('school_name') or 'Unknown'
                    info = _school_info_for_name_level(db, name, level_key)
                    rating = float(info['rating']) if info and info.get('rating') is not None else None
                    addr = (info.get('address') or 'N/A') if info else 'N/A'
                    return name, rating, addr

                elem_name, elem_rating, elem_addr = first_zoned(elem_zones, 'elementary')
                mid_name, mid_rating, mid_addr = first_zoned(mid_zones, 'middle')
                high_name, high_rating, high_addr = first_zoned(high_zones, 'high')

                ratings = [r for r in [elem_rating, mid_rating, high_rating] if r is not None]
                blended_score = sum(ratings) / len(ratings) if ratings else None

                return jsonify({
                    'zip_code': zip_code,
                    'address': address,
                    'latitude': lat,
                    'longitude': lng,
                    'elementary_school_name': elem_name,
                    'elementary_school_rating': elem_rating,
                    'elementary_school_address': elem_addr,
                    'middle_school_name': mid_name,
                    'middle_school_rating': mid_rating,
                    'middle_school_address': mid_addr,
                    'high_school_name': high_name,
                    'high_school_rating': high_rating,
                    'high_school_address': high_addr,
                    'blended_school_score': blended_score,
                    'school_source': 'zoned',
                })

        # STEP 2: Fallback to distance-based nearest schools
        fallback_reason = 'point_not_in_any_zone' if zones else 'no_zones_loaded'

        search_radius = 5.0 / 69.0
        query_params = {
            'lat': lat,
            'lng': lng,
            'lat_min': lat - search_radius,
            'lat_max': lat + search_radius,
            'lng_min': lng - search_radius,
            'lng_max': lng + search_radius
        }

        elementary_query = text("""
            SELECT elementary_school_name, elementary_school_rating, elementary_school_address
            FROM school_data
            WHERE elementary_school_rating IS NOT NULL
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
            ORDER BY 3959 * acos(
                cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) +
                sin(radians(:lat)) * sin(radians(latitude))
            )
            LIMIT 1
        """)
        middle_query = text("""
            SELECT middle_school_name, middle_school_rating, middle_school_address
            FROM school_data
            WHERE middle_school_rating IS NOT NULL
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
            ORDER BY 3959 * acos(
                cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) +
                sin(radians(:lat)) * sin(radians(latitude))
            )
            LIMIT 1
        """)
        high_query = text("""
            SELECT high_school_name, high_school_rating, high_school_address
            FROM school_data
            WHERE high_school_rating IS NOT NULL
              AND latitude BETWEEN :lat_min AND :lat_max
              AND longitude BETWEEN :lng_min AND :lng_max
            ORDER BY 3959 * acos(
                cos(radians(:lat)) * cos(radians(latitude)) * cos(radians(longitude) - radians(:lng)) +
                sin(radians(:lat)) * sin(radians(latitude))
            )
            LIMIT 1
        """)

        elem_result = db.execute(elementary_query, query_params).fetchone()
        mid_result = db.execute(middle_query, query_params).fetchone()
        high_result = db.execute(high_query, query_params).fetchone()

        elementary_name = elem_result[0] if elem_result else None
        elementary_rating = elem_result[1] if elem_result else None
        elementary_addr = elem_result[2] if elem_result else None
        middle_name = mid_result[0] if mid_result else None
        middle_rating = mid_result[1] if mid_result else None
        middle_addr = mid_result[2] if mid_result else None
        high_name = high_result[0] if high_result else None
        high_rating = high_result[1] if high_result else None
        high_addr = high_result[2] if high_result else None

        # Blended score
        ratings = [r for r in [elementary_rating, middle_rating, high_rating] if r is not None]
        blended_score = sum(ratings) / len(ratings) if ratings else None

        # Zip from address
        zip_code = None
        if address:
            import re
            zip_match = re.search(r'\b\d{5}(?:-\d{4})?\b', address)
            if zip_match:
                zip_code = zip_match.group(0)

        result = {
            'zip_code': zip_code,
            'address': address,
            'latitude': lat,
            'longitude': lng,
            'elementary_school_name': elementary_name,
            'elementary_school_rating': elementary_rating,
            'elementary_school_address': elementary_addr,
            'middle_school_name': middle_name,
            'middle_school_rating': middle_rating,
            'middle_school_address': middle_addr,
            'high_school_name': high_name,
            'high_school_rating': high_rating,
            'high_school_address': high_addr,
            'blended_school_score': blended_score,
            'school_source': 'distance_fallback',
            'fallback_reason': fallback_reason,
        }
        return jsonify(result)
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in get_schools_by_address: {e}")
        print(error_trace)
        return jsonify({
            'error': str(e),
            'message': 'Error fetching school data',
            'traceback': error_trace
        }), 500


@api.route('/schools/address/all-zoned', methods=['GET'])
def get_all_zoned_schools_for_address():
    """
    Return ALL NCES attendance zones that contain the given address (point-in-polygon).
    Use for dropdown/export: list every school the address is zoned for (NC/SC only).
    """
    try:
        import requests
        from config.config import Config

        address = request.args.get('address')
        lat = request.args.get('lat', type=float)
        lng = request.args.get('lng', type=float)

        if lat is None or lng is None:
            if not address:
                return jsonify({'error': 'Provide address= or lat= and lng='}), 400
            geocode_url = 'https://maps.googleapis.com/maps/api/geocode/json'
            params = {'address': address, 'key': Config.GOOGLE_MAPS_API_KEY, 'components': 'country:US'}
            response = requests.get(geocode_url, params=params, timeout=10)
            data = response.json()
            if data.get('status') != 'OK' or not data.get('results'):
                return jsonify({'error': 'Could not geocode address'}), 400
            loc = data['results'][0]['geometry']['location']
            lat, lng = loc['lat'], loc['lng']

        db: Session = next(get_db())
        zones = db.query(AttendanceZone).filter(
            or_(AttendanceZone.state == 'NC', AttendanceZone.state == 'SC')
        ).all()
        if not zones:
            return jsonify({
                'address': address,
                'latitude': lat,
                'longitude': lng,
                'elementary': [], 'middle': [], 'high': [],
                'message': 'No NCES attendance zones loaded (NC/SC only).'
            })

        zones_list = [z.to_dict() for z in zones]
        by_level = find_all_zoned_schools(lat, lng, zones_list)

        def to_summary(zone_list):
            return [{'school_name': z.get('school_name'), 'school_level': z.get('school_level'),
                     'school_district': z.get('school_district'), 'state': z.get('state')} for z in zone_list]

        return jsonify({
            'address': address,
            'latitude': lat,
            'longitude': lng,
            'elementary': to_summary(by_level.get('elementary', [])),
            'middle': to_summary(by_level.get('middle', [])),
            'high': to_summary(by_level.get('high', [])),
        })
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500


def _rating_for_school(db: Session, school_name: str, level: str) -> Optional[float]:
    """Look up school rating from school_data by name and level. Returns rating 1-10 or None."""
    if not school_name or not level:
        return None
    level = level.lower()
    if level == 'elementary':
        row = db.query(SchoolData).filter(
            SchoolData.elementary_school_name.ilike(f'%{school_name}%'),
            SchoolData.elementary_school_rating.isnot(None)
        ).first()
        return float(row.elementary_school_rating) if row else None
    if level == 'middle':
        row = db.query(SchoolData).filter(
            SchoolData.middle_school_name.ilike(f'%{school_name}%'),
            SchoolData.middle_school_rating.isnot(None)
        ).first()
        return float(row.middle_school_rating) if row else None
    if level == 'high':
        row = db.query(SchoolData).filter(
            SchoolData.high_school_name.ilike(f'%{school_name}%'),
            SchoolData.high_school_rating.isnot(None)
        ).first()
        return float(row.high_school_rating) if row else None
    return None


def _school_info_for_name_level(
    db: Session, school_name: str, level: str
) -> Optional[Dict]:
    """
    Look up school_data row by name and level. Returns dict with name, rating, address
    (or None for each). Used so UI can show zoned school name even when rating is missing.
    """
    if not school_name or not level:
        return None
    level = level.lower()
    if level == 'elementary':
        row = db.query(SchoolData).filter(
            SchoolData.elementary_school_name.ilike(f'%{school_name}%')
        ).first()
        if not row:
            return None
        return {
            'name': row.elementary_school_name,
            'rating': float(row.elementary_school_rating) if row.elementary_school_rating is not None else None,
            'address': row.elementary_school_address,
        }
    if level == 'middle':
        row = db.query(SchoolData).filter(
            SchoolData.middle_school_name.ilike(f'%{school_name}%')
        ).first()
        if not row:
            return None
        return {
            'name': row.middle_school_name,
            'rating': float(row.middle_school_rating) if row.middle_school_rating is not None else None,
            'address': row.middle_school_address,
        }
    if level == 'high':
        row = db.query(SchoolData).filter(
            SchoolData.high_school_name.ilike(f'%{school_name}%')
        ).first()
        if not row:
            return None
        return {
            'name': row.high_school_name,
            'rating': float(row.high_school_rating) if row.high_school_rating is not None else None,
            'address': row.high_school_address,
        }
    return None


@api.route('/zips/<zip_code>/school-zones', methods=['GET'])
def get_school_zones_by_zip(zip_code: str):
    """
    For a zip code: list school districts that touch the zip, schools per district,
    district strength (avg rating), and GeoJSON geometry for each district's slice of the zip.
    NC/SC only (attendance zones). Requires zip boundary in data/zip_boundaries/{zip}.geojson.
    """
    try:
        zip_polygon = load_zip_polygon(zip_code)
        if zip_polygon is None:
            return jsonify({
                'error': 'Zip boundary not found',
                'message': f'No boundary for zip {zip_code}. Run: python scripts/download_accurate_boundaries.py --zip-codes {zip_code}'
            }), 404
        db: Session = next(get_db())
        zones = db.query(AttendanceZone).filter(
            or_(AttendanceZone.state == 'NC', AttendanceZone.state == 'SC')
        ).all()
        if not zones:
            return jsonify({
                'zip_code': zip_code,
                'district_count': 0,
                'districts': [],
                'message': 'No NCES attendance zones loaded (NC/SC only).'
            })

        zones_list = [z.to_dict() for z in zones]
        intersecting, diag = zones_intersecting_zip_diagnostic(zip_polygon, zones_list)
        if not intersecting:
            return jsonify({
                'zip_code': zip_code,
                'district_count': 0,
                'districts': [],
                'message': 'No attendance zones intersect this zip (NC/SC data only).',
                'debug': diag,
            })

        by_level = request.args.get('by_level', '').lower() in ('1', 'true', 'yes')
        if by_level:
            cache_key = zip_code
            if cache_key in _school_zones_cache:
                return jsonify(_school_zones_cache[cache_key])
            by_level_out = {'elementary': [], 'middle': [], 'high': []}
            LEVEL_COLORS = {'elementary': '#2E7D32', 'middle': '#1565C0', 'high': '#C62828'}
            def _norm_level(s):
                s = (s or 'unknown').lower()
                if 'elem' in s or s == 'elementary': return 'elementary'
                if 'mid' in s or s == 'middle': return 'middle'
                if 'high' in s: return 'high'
                return s
            for level_key in ('elementary', 'middle', 'high'):
                level_zones_raw = [z for z in intersecting if _norm_level(z.get('school_level')) == level_key]
                seen = set()
                level_zones = []
                for z in level_zones_raw:
                    key = (str(z.get('school_name') or '').strip().lower(), level_key)
                    if key in seen:
                        continue
                    seen.add(key)
                    level_zones.append(z)
                for z in level_zones:
                    geometry = zone_geometry_in_zip(zip_polygon, z)
                    if geometry is None:
                        continue
                    name = z.get('school_name') or 'Unknown'
                    rating = None
                    if z.get('canonical_school_id'):
                        school = db.query(School).filter(School.id == z['canonical_school_id']).first()
                        if school:
                            rating = school.rating
                    if rating is None:
                        rating = _rating_for_school(db, name, level_key)
                    by_level_out[level_key].append({
                        'school_name': name,
                        'district_name': z.get('school_district'),
                        'geometry': geometry,
                        'schools': [{'name': name, 'rating': round(rating, 1) if rating is not None else None}],
                        'avg_rating': round(rating, 1) if rating is not None else None,
                        'color': LEVEL_COLORS.get(level_key, '#666666'),
                    })
            resp = {
                'zip_code': zip_code,
                'by_level': True,
                'elementary': by_level_out['elementary'],
                'middle': by_level_out['middle'],
                'high': by_level_out['high'],
            }
            if len(_school_zones_cache) >= _SCHOOL_ZONES_CACHE_MAX:
                _school_zones_cache.pop(next(iter(_school_zones_cache)))
            _school_zones_cache[cache_key] = resp
            return jsonify(resp)

        grouped = group_zones_by_district(intersecting)
        DISTRICT_COLORS = ['#4A90D9', '#50C878', '#E6A23C', '#E07070', '#9B59B6', '#1ABC9C', '#E67E22', '#3498DB']

        districts_out = []
        for i, grp in enumerate(grouped):
            district_id = grp['district_id']
            district_zones = grp['zones']
            schools = []
            ratings = []
            for z in district_zones:
                name = z.get('school_name') or 'Unknown'
                level = (z.get('school_level') or 'unknown').lower()
                rating = _rating_for_school(db, name, level)
                schools.append({'name': name, 'level': level, 'rating': rating})
                if rating is not None:
                    ratings.append(rating)
            avg_rating = sum(ratings) / len(ratings) if ratings else None
            geometry = district_geometry_in_zip(zip_polygon, district_zones)
            color = DISTRICT_COLORS[i % len(DISTRICT_COLORS)]
            districts_out.append({
                'district_id': district_id,
                'district_name': grp['district_name'],
                'schools': schools,
                'avg_rating': round(avg_rating, 1) if avg_rating is not None else None,
                'geometry': geometry,
                'color': color,
            })

        districts_out.sort(key=lambda d: (d['avg_rating'] is None, -(d['avg_rating'] or 0)))

        return jsonify({
            'zip_code': zip_code,
            'district_count': len(districts_out),
            'districts': districts_out,
        })
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500


@api.route('/schools/zip/<zip_code>', methods=['GET'])
def get_schools_by_zip(zip_code: str):
    """Get school ratings summary for a zip code (single row if cached)."""
    try:
        db: Session = next(get_db())

        # Check if we have cached data
        cached = db.query(SchoolData).filter(SchoolData.zip_code == zip_code).first()
        if cached:
            return jsonify(cached.to_dict())

        # If not cached, we need an address to geocode
        return jsonify({
            'error': 'School data not found for this zip code',
            'message': 'Please use /api/schools/address?address=<full_address> to fetch school data'
        }), 404

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@api.route('/schools/zip/<zip_code>/list', methods=['GET'])
def list_schools_by_zip(zip_code: str):
    """List unique schools in a zip code for plotting on map. Returns name, level, address, lat, lng, rating."""
    try:
        db: Session = next(get_db())
        rows = db.execute(text("""
            SELECT DISTINCT ON (LOWER(TRIM(name)), level)
                   name, level, address, latitude, longitude, rating
            FROM (
                SELECT elementary_school_name AS name, 'elementary' AS level,
                       COALESCE(elementary_school_address, address) AS address,
                       latitude, longitude, elementary_school_rating AS rating
                FROM school_data
                WHERE zip_code = :zip AND elementary_school_name IS NOT NULL AND elementary_school_rating IS NOT NULL
                UNION ALL
                SELECT middle_school_name, 'middle',
                       COALESCE(middle_school_address, address),
                       latitude, longitude, middle_school_rating
                FROM school_data
                WHERE zip_code = :zip AND middle_school_name IS NOT NULL AND middle_school_rating IS NOT NULL
                UNION ALL
                SELECT high_school_name, 'high',
                       COALESCE(high_school_address, address),
                       latitude, longitude, high_school_rating
                FROM school_data
                WHERE zip_code = :zip AND high_school_name IS NOT NULL AND high_school_rating IS NOT NULL
            ) sub
            WHERE name IS NOT NULL
            ORDER BY LOWER(TRIM(name)), level, rating DESC
        """), {"zip": zip_code}).fetchall()

        schools = [
            {"name": r[0], "level": r[1], "address": r[2], "latitude": r[3], "longitude": r[4], "rating": float(r[5]) if r[5] else None}
            for r in rows
        ]
        return jsonify({"zip_code": zip_code, "schools": schools})

    except Exception as e:
        return jsonify({'error': str(e)}), 500
